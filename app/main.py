import os
import base64
import time
import signal
import sys
from gpt_communication.gemini_gpt import GenerativeModelHandler
from docx import Document
from docx.shared import Pt

# 🔹 Handler para interrupção do usuário
def handler(signum, frame):
    print("Processamento interrompido pelo usuário")
    sys.exit(1)

signal.signal(signal.SIGINT, handler)

# 🔹 Diretório do script
current_dir = os.path.dirname(os.path.abspath(__file__))

# 🔹 Diretório das imagens
assets_dir = os.path.join(current_dir, "assets")

# 🔹 Caminho do documento Word
output_docx = os.path.join(current_dir, "resumo_analises_imagens.docx")

# 🔹 Criar instância do handler reutilizando a configuração existente
gemini_handler = GenerativeModelHandler("gemini-2.0-flash-exp")

task_prompt = (
    "Descreva em português claro e objetivo qual o passo que a imagem representa na criação de um agente no Copilot Studio. "
    "Se for uma configuração, explique que tipo de parâmetro ou ajuste está sendo realizado (exemplo: configuração de comportamento, adição de frases-gatilho, ou personalização de resposta). "
    "Imagens dos testes começam na imagem custom_agent_14.png. Se for um teste, explique que tipo de interação está sendo realizada (exemplo: teste de resposta, validação de ação, ou simulação de conversa). "
    "Evite descrever textos nas imagens, apenas forneça um resumo sobre se é uma configuração ou teste. "
    "Se houver texto em inglês, traduza-o para português, sendo breve e direto."
)


# 🔹 Abrir documento existente ou criar novo
if os.path.exists(output_docx):
    doc = Document(output_docx)
else:
    doc = Document()
    doc.add_heading('Resumo das Análises das Imagens', level=0)

# 🔹 Listar imagens no diretório assets ordenadas pela data de modificação (mais antiga primeiro)
image_files = [
    f for f in os.listdir(assets_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg'))
]

# 🔹 Ordenar arquivos pela data de modificação (mais antiga primeiro)
image_files.sort(key=lambda x: os.path.getmtime(os.path.join(assets_dir, x)))

# 🔹 Processar cada imagem individualmente
for idx, image_name in enumerate(image_files, start=1):
    img_path = os.path.join(assets_dir, image_name)

    print(f"\n🔄 Processando imagem {idx}/{len(image_files)}: {image_name}")

    # 🔹 Ler o conteúdo da imagem e converter para Base64
    with open(img_path, "rb") as img_file:
        img_data = base64.b64encode(img_file.read()).decode("utf-8")

    try:
        # 🔹 Enviar a imagem e obter resposta
        response_text = gemini_handler.generate_content_with_image_data(img_path, task_prompt)

        # 🔹 Exibir a resposta no terminal
        print(f"✅ Resumo da imagem '{image_name}':\n{response_text}")

        # 🔹 Salvar resultado no documento Word existente
        doc.add_heading(f"Imagem: {image_name}", level=1)
        p = doc.add_paragraph(response_text)
        p.style.font.size = Pt(12)

    except Exception as e:
        print(f"❌ Erro ao processar '{image_name}': {str(e)}")
        doc.add_heading(f"Imagem: {image_name}", level=1)
        doc.add_paragraph(f"Erro ao processar imagem: {str(e)}")

    # 🔹 Aguardar 10 segundos antes da próxima requisição
    if idx < len(image_files):
        print("⏳ Aguardando alguns segundos para próxima requisição...")
        time.sleep(7)

# 🔹 Salvar documento Word atualizado
doc.save(output_docx)
print(f"\n📄 Documento atualizado com sucesso em: {output_docx}")