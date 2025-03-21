import os
from docx import Document
from docx.shared import Pt
from core.config import OUTPUT_DOCX

class DocumentService:
    def __init__(self):
        self.doc = self._load_or_create_document()

    def _load_or_create_document(self):
        if os.path.exists(OUTPUT_DOCX):
            return Document(OUTPUT_DOCX)
        doc = Document()
        doc.add_heading('Resumo das Análises das Imagens', level=0)
        return doc

    def add_image_summary(self, image_name, summary):
        self.doc.add_heading(f"Imagem: {image_name}", level=1)
        p = self.doc.add_paragraph(summary)
        p.style.font.size = Pt(12)

    def save_document(self):
        self.doc.save(OUTPUT_DOCX)